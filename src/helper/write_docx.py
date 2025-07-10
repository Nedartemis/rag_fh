from enum import Enum
from pathlib import Path
from typing import List, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from vars import PATH_TMP

LIGHT_BLUE = "a4c2f4"
LIGHT_GREY = "efefef"


def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


class DocxWriter:

    def __init__(self):
        self.doc = Document()

    def add_text(
        self,
        content: str,
        fontsize: int,
        underline: bool = False,
        alignement=WD_PARAGRAPH_ALIGNMENT.LEFT,
    ):
        text = self.doc.add_paragraph(content)
        text.alignment = alignement
        run = text.runs[0]
        run.font.size = Pt(fontsize)
        run.underline = underline

    def add_newline(self) -> None:
        self.doc.add_paragraph("")

    def add_table(
        self,
        content: pd.DataFrame,
        color_header: str,
        color_cells: str,
        column_widths: Optional[List[float]] = None,
    ) -> None:
        nrows, ncols = content.shape
        table = self.doc.add_table(rows=nrows + 1, cols=ncols)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False

        assert len(column_widths) == ncols

        # header
        for i, text in enumerate(content.columns):
            cell = table.cell(0, i)
            cell.text = text
            set_cell_background(cell, color_header)

        # cells
        for row_idx, (_, row) in enumerate(content.iterrows()):
            for col_idx, text in enumerate(row):
                cell = table.cell(row_idx + 1, col_idx)
                cell.text = text
                set_cell_background(cell, color_cells)

        # width
        if column_widths is None:
            column_widths = [6 / ncols] * ncols

        for idx, width in enumerate(column_widths):
            table.columns[idx].width = Inches(width)

    def save(self, path_doc: Path) -> None:
        self.doc.save(path_doc)


if __name__ == "__main__":
    doc = DocxWriter()

    # --- Titre centré, police 18 ---
    doc.add_text(
        content="Mon titre", fontsize=18, alignement=WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    doc.add_newline()

    # --- Sous-titres, police 14 ---
    doc.add_text(content="Sous-titre 1", fontsize=14, underline=True)

    doc.add_newline()

    # --- Création du tableau ---
    doc.add_table(
        pd.DataFrame(data=[{"1": "toto", "2": "foo"}]),
        color_header=LIGHT_BLUE,
        color_cells=LIGHT_GREY,
    )

    # Enregistrement
    doc.save(PATH_TMP / "mon_document.docx")
