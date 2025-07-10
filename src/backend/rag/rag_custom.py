from pathlib import Path
from typing import List, Optional, Tuple

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from tqdm import tqdm

from backend.ocr import ocr_pdf
from backend.rag.rag_pipeline import RagPipeline
from backend.read_pdf import read_pdf
from vars import PATH_DOCS

CHUNK_SIZE = 512 * 2


def is_scanned_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        text = doc[page_num].get_text()
        if text.strip():  # Found some text
            return False
    return True  # No text found in any page


def read_docx_by_page(filename) -> List[str]:
    doc = Document(filename)
    texts = []

    for para in doc.paragraphs:
        texts.append(para.text)
        print(para.text)
        # for run in para.runs:
        #     if run.break_type == WD_BREAK.PAGE:
        #         # When a page break is found, save current content
        #         pages.append("\n".join(current_page).strip())
        #         current_page = []

    # Append any remaining content after the last page break
    # if current_page:
    #     pages.append("\n".join(current_page).strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)

    text = ["\n".join(texts)]
    return text


def file_to_text(path_file: Path) -> List[dict]:
    if path_file.suffix == ".txt":
        with open(path_file, "r") as fr:
            text = fr.read()
        return [{"pathfile": path_file, "text": text, "page": None}]
    elif path_file.suffix == ".docx":
        texts = read_docx_by_page(path_file)
        return [
            {"pathfile": path_file, "text": text, "page": None}
            for page, text in enumerate(texts, start=1)
        ]
    elif path_file.suffix == ".pdf":
        if is_scanned_pdf(path_file):
            text = "\n".join(ocr_pdf(path_file))
        else:
            text = "\n".join(read_pdf(path_file))
        return [{"pathfile": path_file, "text": text, "page": None}]

    print(f"Extension '{path_file.suffix}' not handled.")
    return []


class RagCustom(RagPipeline):

    def __init__(self):
        super().__init__()

    def ask(self, question: str, path_files: List[str]) -> str:

        # load from cache

        # convert into texts
        data = [
            e
            for path_file in tqdm(path_files, desc="Reading documents")
            for e in file_to_text(path_file)
        ]
        df_text = pd.DataFrame(data)

        # text into chunks
        data = []
        for pathfile in path_files:
            df = df_text[df_text["pathfile"] == pathfile]
            idx = 0

            for _, row in df.iterrows():
                text, page = row["text"], row["page"]
                idx = 0
                while idx <= len(text):
                    chunk = text[idx : idx + CHUNK_SIZE]
                    data.append({"pathfile": pathfile, "chunk": chunk, "page": page})
                    idx += CHUNK_SIZE // 2

        df: pd.DataFrame = pd.DataFrame(data)
        print(f"{len(df)} chunks")

        # encoding
        embeddings = self.retriever.encode(df["chunk"])

        # save df and embeddings

        # ask
        return self._ask(question, df, embeddings)

    def format_chunks(self, df) -> List[str]:
        return [
            f"Filename : {row['pathfile']} :\n\n{row['chunk']}"
            for _, row in df.iterrows()
        ]

    def get_instructions(self):
        return ["Bien mettre les sources de l'information (nom du fichier)"]

    def get_n(self) -> int:
        return 5


if __name__ == "__main__":
    rag = RagCustom()
    rag.ask(
        question="Quelles sont les noms des différents lots ?",
        path_files=[
            # PATH_DOCS / "assignation.txt",
            # PATH_DOCS / "Estimation Frais et honoraires.docx",
            PATH_DOCS / "intégrale CR chantier Saint Amand.pdf_page_1-7.pdf",
            # PATH_DOCS
            # / "TJ DIEPPE - MASSERE - Ordonnance Rectificative du 15 01 2025.pdf",
        ],
    )
